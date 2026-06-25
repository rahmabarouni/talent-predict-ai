"""CV Text Extraction Router."""

import logging
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pathlib import Path
import PyPDF2
from docx import Document
import io

router = APIRouter(prefix="/cv", tags=["CV Extraction"])
logger = logging.getLogger(__name__)

# Allowed file types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.extract_text()
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num + 1}: {str(e)}")
                continue
        return text
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_txt_text(file_bytes: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")

@router.post("/extract")
async def extract_cv(file: UploadFile = File(...)) -> JSONResponse:
    """Extract text from CV file."""
    try:
        # Validate file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type '{file_ext}' not supported. Allowed: {ALLOWED_EXTENSIONS}"
            )
        
        # Read file bytes
        file_bytes = await file.read()
        
        # Check file size
        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed ({MAX_FILE_SIZE / 1024 / 1024}MB)"
            )
        
        # Extract text based on file type
        logger.info(f"Extracting text from: {file.filename}")
        
        if file_ext == ".pdf":
            extracted_text = extract_pdf_text(file_bytes)
        elif file_ext in {".docx", ".doc"}:
            extracted_text = extract_docx_text(file_bytes)
        elif file_ext == ".txt":
            extracted_text = extract_txt_text(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unknown file type")
        
        # Clean and validate extracted text
        extracted_text = extracted_text.strip()
        if not extracted_text:
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the file"
            )
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters")
        
        return JSONResponse({
            "success": True,
            "extracted_text": extracted_text,
            "filename": file.filename,
            "file_type": file_ext,
            "size": len(file_bytes),
            "text_length": len(extracted_text)
        })
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting CV: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text: {str(e)}"
        )

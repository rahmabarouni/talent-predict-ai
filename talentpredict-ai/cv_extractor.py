"""
CV Text Extraction Service
Extracts text from PDF and Word documents for skill analysis
"""

import os
import logging
import io
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
import PyPDF2
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="CV Text Extractor",
    description="Extract text from CV files (PDF, DOCX)",
    version="1.0.0"
)

# Allowed file types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def extract_pdf_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF file with robust error handling"""
    try:
        if not file_bytes:
            logger.error(f"Empty PDF file: {filename}")
            raise HTTPException(
                status_code=422,
                detail={"error": "PDF is empty", "code": "PDF_EMPTY"}
            )

        stream = io.BytesIO(file_bytes)
        try:
            pdf_reader = PyPDF2.PdfReader(stream)
        except Exception as e:
            logger.error(f"Corrupted PDF file: {filename} - {str(e)}")
            raise HTTPException(
                status_code=422,
                detail={"error": "PDF could not be read", "code": "PDF_CORRUPT"}
            )

        if pdf_reader.is_encrypted:
            logger.error(f"Password protected PDF: {filename}")
            raise HTTPException(
                status_code=422,
                detail={"error": "PDF is password protected", "code": "PDF_ENCRYPTED"}
            )

        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num + 1} from {filename}: {str(e)}")
                continue
        
        if not text.strip():
             logger.warning(f"No text extracted from PDF: {filename}")
             # We don't throw here, the main block will handle empty string check

        return text
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected PDF extraction error for {filename}: {str(e)}")
        raise HTTPException(
            status_code=422,
            detail={"error": "PDF could not be read", "code": "PDF_CORRUPT"}
        )


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


def extract_txt_text(file_bytes: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


@app.post("/extract")
async def extract_cv(file: UploadFile = File(...)) -> JSONResponse:
    """
    Extract text from CV file
    
    Supported formats: PDF, DOCX, DOC, TXT
    
    Returns:
        - extracted_text: Full text from CV
        - filename: Original filename
        - file_type: Type of file processed
        - size: File size in bytes
    """
    try:
        # Validate file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type '{file_ext}' not supported. Allowed: {ALLOWED_EXTENSIONS}"
            )
        
        # Read file bytes
        file_bytes = await file.read()
        
        # Check file size
        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed ({MAX_FILE_SIZE / 1024 / 1024}MB)"
            )
        
        # Extract text based on file type
        logger.info(f"Extracting text from: {file.filename}")
        
        if file_ext == ".pdf":
            extracted_text = extract_pdf_text(file_bytes, file.filename)
        elif file_ext in {".docx", ".doc"}:
            extracted_text = extract_docx_text(file_bytes)
        elif file_ext == ".txt":
            extracted_text = extract_txt_text(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unknown file type")
        
        # Clean and validate extracted text
        extracted_text = extracted_text.strip()
        if not extracted_text:
            logger.error(f"Extraction resulted in empty text for: {file.filename}")
            raise HTTPException(
                status_code=422,
                detail={"error": "PDF is empty", "code": "PDF_EMPTY"}
            )
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters")
        
        return JSONResponse({
            "success": True,
            "extracted_text": extracted_text,
            "filename": file.filename,
            "file_type": file_ext,
            "size": len(file_bytes),
            "text_length": len(extracted_text)
        })
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting CV {file.filename}: {str(e)}")
        raise HTTPException(
            status_code=422,
            detail={"error": "PDF could not be read", "code": "PDF_CORRUPT"}
        )


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "CV Text Extractor"}


@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "service": "CV Text Extractor",
        "version": "1.0.0",
        "endpoints": {
            "POST /extract": "Extract text from CV file",
            "GET /health": "Health check",
            "GET /": "This message"
        }
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=9000)
